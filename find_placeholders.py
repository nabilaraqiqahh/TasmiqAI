import docx

def find_placeholders():
    doc = docx.Document("PSM Report Fixed.docx")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "diagram here" in text.lower():
            # Print the paragraph before it to see the figure title
            title = doc.paragraphs[i-1].text if i > 0 else ""
            if not title.strip() and i > 1:
                title = doc.paragraphs[i-2].text
            print(f"Placeholder found at {i}: {title.strip()}")

if __name__ == "__main__":
    find_placeholders()

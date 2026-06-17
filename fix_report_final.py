import docx

def replace_heading(para, new_text):
    print(f"Fixing: {para.text} -> {new_text}")
    # Remove auto-numbering if it exists so we can manually add the number
    if para._p.pPr and para._p.pPr.numPr:
        para._p.pPr.remove(para._p.pPr.numPr)
    
    # Replace text while keeping the format of the first run
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = new_text

def fix_report():
    doc = docx.Document("PSM Report.docx")
    
    # First pass: Unique headings
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        # Chapter 4 unique replacements
        if text == "4.3 System Architecture": replace_heading(para, "4.2.1 System Architecture")
        elif text == "4.4 User Interface Design": replace_heading(para, "4.2.2 User Interface Design")
        elif text == "4.4.1 Navigation Design": replace_heading(para, "4.2.2.1 Navigation Design")
        elif text == "4.4.2 Input Design": replace_heading(para, "4.2.2.2 Input Design")
        elif text == "4.4.3 Output Design": replace_heading(para, "4.2.2.3 Output Design")
        elif text == "4.5 Database Design": replace_heading(para, "4.2.3 Database Design")
        elif text == "4.5.1 Conceptual and Logical Database Design": replace_heading(para, "4.2.3.1 Conceptual and Logical Database Design")
        elif text == "4.5.2 Entity Relationship Diagram (ERD)": replace_heading(para, "4.2.3.2 Entity Relationship Diagram (ERD)")
        elif text == "4.5.3 Data Dictionary": replace_heading(para, "4.2.3.3 Data Dictionary")
        elif text == "4.5.4 Normalization": replace_heading(para, "4.2.3.4 Normalization")
        elif text == "4.6 Detailed Design": replace_heading(para, "4.3 Detailed Design")
        elif text == "4.6.1 Software Design (Backend - FastAPI)": replace_heading(para, "4.3.1 Software Design (Backend - FastAPI)")
        elif text == "4.6.2 Software Design (AI Engine)": replace_heading(para, "4.3.2 Software Design (AI Engine)")
        elif text == "4.6.3 Software Design (Mobile Application)": replace_heading(para, "4.3.3 Software Design (Mobile Application)")
        elif text == "4.6.4 Software Design (Teacher Portal)": replace_heading(para, "4.3.4 Software Design (Teacher Portal)")
        elif text == "4.7 Physical Database Design": replace_heading(para, "4.3.5 Physical Database Design")
        elif text == "4.8 Conclusion": replace_heading(para, "4.4 Conclusion")

        # Chapter 6 unique replacements
        elif text == "Test Plan" and para.style.name.startswith("Heading"): replace_heading(para, "6.2 Test Plan")
        elif text == "Test Organization" and para.style.name.startswith("Heading"): replace_heading(para, "6.2.1 Test Organization")
        elif text == "Test Environment" and para.style.name.startswith("Heading"): replace_heading(para, "6.2.2 Test Environment")
        elif text == "Classes of tests" and para.style.name.startswith("Heading"): replace_heading(para, "6.3.1 Classes of tests")
        elif text == "Test Design" and para.style.name.startswith("Heading"): replace_heading(para, "6.4 Test Design")
        elif text == "Test Description" and para.style.name.startswith("Heading"): replace_heading(para, "6.4.1 Test Description")
        elif text == "Test Data" and para.style.name.startswith("Heading"): replace_heading(para, "6.4.2 Test Data")
        elif text == "Test Results and Analysis" and para.style.name.startswith("Heading"): replace_heading(para, "6.5 Test Results and Analysis")

    # Second pass: Non-unique headings that depend on chapter context
    chapter = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        
        # Only update chapter if it's a Heading and starts with CHAPTER
        if text.upper().startswith("CHAPTER") and para.style.name.startswith("Heading"):
            if "1" in text: chapter = 1
            elif "2" in text: chapter = 2
            elif "3" in text: chapter = 3
            elif "4" in text: chapter = 4
            elif "5" in text: chapter = 5
            elif "6" in text: chapter = 6
            elif "7" in text: chapter = 7
            
        if chapter == 2 and text == "Conclusion" and para.style.name.startswith("Heading"):
            replace_heading(para, "2.6 Conclusion")
        elif chapter == 3 and para.style.name.startswith("Heading"):
            if text == "Introduction": replace_heading(para, "3.1 Introduction")
            elif text == "Problem Analysis": replace_heading(para, "3.2 Problem Analysis")
            elif text == "Requirement analysis": replace_heading(para, "3.3 Requirement analysis")
            elif text == "Conclusion": replace_heading(para, "3.4 Conclusion")
        elif chapter == 6 and text == "Conclusion" and para.style.name.startswith("Heading"):
            replace_heading(para, "6.6 Conclusion")
            
    doc.save("PSM Report Fixed.docx")
    print("Successfully saved to 'PSM Report Fixed.docx'")

if __name__ == "__main__":
    fix_report()

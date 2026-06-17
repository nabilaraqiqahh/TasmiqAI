import docx
from docx.shared import Inches
import urllib.request
import zlib
import base64

diagrams = {
    "Figure 4.2": """graph TD
    subgraph Client Tier
        Mobile[Student Mobile App]
        Portal[Teacher Web Portal]
    end
    subgraph Application Tier
        API[FastAPI Backend]
        AI[AI Engine Pipeline]
    end
    subgraph Data Tier
        DB[(Supabase Database)]
        Storage[(Audio Storage)]
    end
    Mobile -- HTTP POST --> API
    Portal -- HTTP GET/POST --> API
    API <--> AI
    API --> DB
    API --> Storage
""",
    "Figure 4.3": """stateDiagram-v2
    [*] --> Splash
    Splash --> Login
    Login --> Dashboard
    Dashboard --> Recitation
    Dashboard --> LearningMode
    Dashboard --> History
    Recitation --> SurahSelection
    SurahSelection --> Recording
    Recording --> Feedback
    Feedback --> Dashboard
""",
    "Figure 4.4": """stateDiagram-v2
    [*] --> Login
    Login --> Dashboard
    Dashboard --> StudentsList
    Dashboard --> PendingRecitations
    Dashboard --> Analytics
    StudentsList --> StudentDetail
    PendingRecitations --> RecitationReview
    RecitationReview --> FeedbackSubmission
""",
    "Figure 4.5": """erDiagram
    USERS ||--o{ RECITATIONS : "submits"
    RECITATIONS ||--|| ASSESSMENTS : "has"
    TEACHERS ||--o{ CLASSES : "supervises"
    CLASSES ||--o{ CLASSMEMBERS : "has"
    USERS ||--o{ CLASSMEMBERS : "belongs to"
    RECITATIONS ||--o{ FEEDBACK : "receives"
    TEACHERS ||--o{ FEEDBACK : "provides"
""",
    "Figure 4.6": """graph TD
    Main[main.py - Entry Point] --> Routers
    Routers --> AuthRouter
    Routers --> RecitationRouter
    Routers --> TeacherRouter
    Routers --> Services
    Services --> AuthService
    Services --> AudioService
    Services --> AIService
    Services --> DataManagers
    DataManagers --> DBClient
    DataManagers --> StorageClient
""",
    "Figure 4.7": """graph LR
    Audio([Raw Audio]) --> Pre[Audio Preprocessing]
    Pre --> ASR[Whisper Speech Recognition]
    ASR --> Text[Transcribed Text]
    Text --> Match[Text Comparison]
    Match --> Tajwid[Tajwid Rule Checking]
    Tajwid --> Score[Scoring & Feedback]
    Score --> Result([Assessment Result])
"""
}

def get_kroki_url(text):
    compressed = zlib.compress(text.encode('utf-8'), 9)
    b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/mermaid/png/{b64}"

def insert_diagrams():
    doc = docx.Document("PSM Report Fixed.docx")
    
    image_paths = {}
    for key, mermaid_code in diagrams.items():
        url = get_kroki_url(mermaid_code)
        path = f"{key.replace(' ', '_').replace('.', '_')}.png"
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
                out_file.write(response.read())
            image_paths[key] = path
            print(f"Downloaded {key}")
        except Exception as e:
            print(f"Failed to download {key}: {e}")

    inserted = set()
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if 'TOC' in para.style.name:
            continue
            
        for key in image_paths:
            if text.startswith(key) and key not in inserted:
                inserted.add(key)
                
                found_placeholder = False
                for j in range(1, 4):
                    if i + j < len(doc.paragraphs):
                        next_para = doc.paragraphs[i+j]
                        if "*diagram here*" in next_para.text.lower():
                            next_para.text = ""
                            run = next_para.add_run()
                            run.add_picture(image_paths[key], width=Inches(6.0))
                            found_placeholder = True
                            print(f"Inserted {key} over placeholder")
                            break
                
                if not found_placeholder:
                    if i + 1 < len(doc.paragraphs):
                        new_p = doc.paragraphs[i+1].insert_paragraph_before()
                        run = new_p.add_run()
                        run.add_picture(image_paths[key], width=Inches(6.0))
                        print(f"Inserted {key} after title")

    doc.save("PSM Report Fixed With Diagrams.docx")
    print("Saved document with diagrams!")

if __name__ == "__main__":
    insert_diagrams()

import docx
from docx.shared import Inches
import urllib.request
import zlib
import base64
import os

diagrams = {
    "Figure 4.2": """graph TD
    subgraph Client Tier
        Mobile[Student Mobile App]
        Portal[Teacher Web Portal]
    end
    subgraph Application Tier
        API[FastAPI Backend]
        AI[AI Engine Pipeline]
    end
    subgraph Data Tier
        DB[(Supabase Database)]
        Storage[(Audio Storage)]
    end
    Mobile -- HTTP POST --> API
    Portal -- HTTP GET/POST --> API
    API <--> AI
    API --> DB
    API --> Storage
""",
    "Figure 4.3": """stateDiagram-v2
    [*] --> Splash
    Splash --> Login
    Login --> Dashboard
    Dashboard --> Recitation
    Dashboard --> LearningMode
    Dashboard --> History
    Recitation --> SurahSelection
    SurahSelection --> Recording
    Recording --> Feedback
    Feedback --> Dashboard
""",
    "Figure 4.4": """stateDiagram-v2
    [*] --> Login
    Login --> Dashboard
    Dashboard --> StudentsList
    Dashboard --> PendingRecitations
    Dashboard --> Analytics
    StudentsList --> StudentDetail
    PendingRecitations --> RecitationReview
    RecitationReview --> FeedbackSubmission
""",
    "Figure 4.5": """erDiagram
    USER ||--o{ RECITATION : "submits"
    RECITATION ||--|| ASSESSMENT : "has"
    TEACHER ||--o{ CLASS : "supervises"
    CLASS }o--o{ USER : "contains"
    RECITATION ||--o{ FEEDBACK : "receives"
    TEACHER ||--o{ FEEDBACK : "provides"
""",
    "Figure 4.6": """graph TD
    Main[main.py - Entry Point] --> Routers
    Routers --> AuthRouter
    Routers --> RecitationRouter
    Routers --> TeacherRouter
    Routers --> Services
    Services --> AuthService
    Services --> AudioService
    Services --> AIService
    Services --> DataManagers
    DataManagers --> DBClient
    DataManagers --> StorageClient
""",
    "Figure 4.7": """graph LR
    Audio([Raw Audio]) --> Pre[Audio Preprocessing]
    Pre --> ASR[Whisper Speech Recognition]
    ASR --> Text[Transcribed Text]
    Text --> Match[Text Comparison]
    Match --> Tajwid[Tajwid Rule Checking]
    Tajwid --> Score[Scoring & Feedback]
    Score --> Result([Assessment Result])
"""
}

def get_kroki_url(text):
    compressed = zlib.compress(text.encode('utf-8'), 9)
    b64 = base64.urlsafe_b64encode(compressed).decode('ascii')
    return f"https://kroki.io/mermaid/png/{b64}"

def insert_diagrams():
    doc = docx.Document("PSM Report Fixed.docx")
    
    # Generate and download images
    image_paths = {}
    for key, mermaid_code in diagrams.items():
        url = get_kroki_url(mermaid_code)
        path = f"{key.replace(' ', '_').replace('.', '_')}.png"
        try:
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req) as response, open(path, 'wb') as out_file:
                out_file.write(response.read())
            image_paths[key] = path
            print(f"Downloaded {key}")
        except Exception as e:
            print(f"Failed to download {key}: {e}")

    # Now find insertion points
    # We will look for "Figure X.Y:" and then if the next paragraph is "*diagram here*", replace it.
    # Otherwise insert after.
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for key in image_paths:
            if text.startswith(key):
                print(f"Found insertion point for {key} at para {i}")
                
                # Check next paragraph
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i+1]
                    if "*diagram here*" in next_para.text.lower() or next_para.text.strip() == "":
                        print(f"Replacing placeholder for {key}")
                        next_para.text = ""
                        run = next_para.add_run()
                        run.add_picture(image_paths[key], width=Inches(6.0))
                        break
                
                # If no placeholder was found, we need to insert a new paragraph.
                # However, python-docx doesn't easily allow inserting paragraphs in the middle.
                # But it has insert_paragraph_before()
                print(f"No explicit placeholder for {key}, inserting after")
                if i + 1 < len(doc.paragraphs):
                    new_p = doc.paragraphs[i+1].insert_paragraph_before()
                    run = new_p.add_run()
                    run.add_picture(image_paths[key], width=Inches(6.0))
                break

    doc.save("PSM Report Fixed.docx")
    print("Saved document with diagrams!")

if __name__ == "__main__":
    insert_diagrams()
